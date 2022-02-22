import os
import re
import time
import urllib.request

# docx imports
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx2pdf import convert


from selenium import webdriver
from selenium.webdriver import ActionChains
from selenium.webdriver.common.keys import Keys

driver = webdriver.Chrome("chromedriver.exe")
driver.set_window_size(1920, 1080)

time.sleep(5)

with open('links.txt', 'r') as links_file:
    test_links = links_file.read().split("\n")

for test_link in test_links:
    driver.get(test_link)
    time.sleep(2)

    test_title = driver.find_element_by_class_name("Header_Title__3mXad")
    test_title_text = re.findall(
        r'<p>(.*)</p>', test_title.get_attribute('innerHTML'))[0]
    try:
        os.mkdir(test_title_text)
    except:
        pass

    buttons = driver.find_elements_by_class_name("Footer_MoveToQuestionBig__1-jSA")
    question = 1
    elements = driver.find_elements_by_class_name(
        "Question_QuestionContainer__2Elg3")

    questions_image_furls = []
    answers_image_furls = []

    while(True):
        # question images
        text = elements[0].get_attribute('innerHTML')
        links = re.findall(r'src="([^"]*)"', text)
        i = 1
        question_image_furls = []
        for link in links:
            urllib.request.urlretrieve(
                link, test_title_text + "/Question " + str(question) + "." + str(i) + ".png")
            question_image_furls.append(
                test_title_text + "/Question " + str(question) + "." + str(i) + ".png")
            i += 1
        questions_image_furls.append(question_image_furls)

        # answer images
        text = elements[1].get_attribute('innerHTML')
        links = re.findall(r'src="([^"]*)"', text)
        i = 1
        answer_image_furls = []
        if links:
            for link in links:
                urllib.request.urlretrieve(
                    link, test_title_text + "/Answer " + str(question) + "." + str(i) + ".png")
                answer_image_furls.append(
                    test_title_text + "/Answer " + str(question) + "." + str(i) + ".png")
                i += 1
        else:
            print("Text solution or No solution detected. Please recheck screenshot")
            answer = driver.find_element_by_class_name("Options_Container__2YfYv")
            ActionChains(driver).move_to_element(answer).click().send_keys(Keys.ARROW_DOWN).send_keys(Keys.ARROW_DOWN).perform()
            time.sleep(1)
            png = answer.screenshot_as_png
            with open(test_title_text + "/Answer " + str(question) + "." + str(i) + ".png", 'wb') as f:
                f.write(png)
                answer_image_furls.append(
                    test_title_text + "/Answer " + str(question) + "." + str(i) + ".png")
            i += 1
            if elements[1].size['height']:
                ActionChains(driver).move_to_element(elements[1]).perform()
                time.sleep(1)
                png = elements[1].screenshot_as_png
                with open(test_title_text + "/Answer " + str(question) + "." + str(i) + ".png", 'wb') as f:
                    f.write(png)
                    answer_image_furls.append(
                        test_title_text + "/Answer " + str(question) + "." + str(i) + ".png")
        answers_image_furls.append(answer_image_furls)

        print("Question " + str(question) + " done")
        question += 1

        if buttons[1].is_enabled() == False:
            break

        buttons[1].click()
        # time.sleep(1)

    driver.close()

    # docx prep start
    document = Document()
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('heading1', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(20)
    obj_font.name = 'Arial'
    obj_font.bold = True

    obj2_charstyle = obj_styles.add_style('heading2', WD_STYLE_TYPE.PARAGRAPH)
    obj2_font = obj2_charstyle.font
    obj2_font.size = Pt(20)
    obj2_font.name = 'Arial'
    obj2_font.bold = False

    obj2_charstyle = obj_styles.add_style('normal', WD_STYLE_TYPE.PARAGRAPH)
    obj2_font = obj2_charstyle.font
    obj2_font.size = Pt(13)
    obj2_font.name = 'Arial'
    obj2_font.bold = False

    # add test name as title to  doc
    title = document.add_paragraph(test_title_text, style='heading1')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('Questions:', style='heading2')

    for q_no, q_furls in enumerate(questions_image_furls):
        # add question no to doc
        document.add_paragraph('Q' + str(q_no+1)+':', style='normal')
        for q_furl in q_furls:
            try:
                document.add_picture(q_furl, width=Inches(5.5))
            except:
                document.add_paragraph('some err', style='Intense Quote')
        # document.add_page_break()

    document.add_page_break()
    document.add_paragraph('Answers:', style='heading2')

    for ans_no, ans_furls in enumerate(answers_image_furls):
        # add ans no to doc
        document.add_paragraph('A' + str(ans_no+1) + ':', style='normal')
        for ans_furl in ans_furls:
            try:
                document.add_picture(ans_furl, width=Inches(5.5))
            except:
                document.add_paragraph('some err', style='Intense Quote')
        # document.add_page_break()

    document.save(test_title_text+"/"+test_title_text+'.docx')
    print("DOCX Generated")

    convert(test_title_text+"/"+test_title_text+'.docx')
    print("PDF Generated")